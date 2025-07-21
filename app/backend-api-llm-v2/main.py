import os
import re
from typing import List
from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from dotenv import load_dotenv
import requests
import json
import logging
import os
import uuid
import shutil
from pathlib import Path
from datetime import datetime
import fitz  # PyMuPDF for PDF extraction
import docx  # python-docx for DOCX files
from typing import Optional, Dict, Any
from fastapi import UploadFile, File, Form
from fastapi.responses import JSONResponse
import hashlib
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter

from langchain.embeddings import OpenAIEmbeddings


from langchain.vectorstores import FAISS
from langchain.schema import Document

logger = logging.getLogger("uvicorn.error")

# Helper function to correctly format EPC scores
def format_epc_score(epc_score):
    """Convert EPC score format correctly, handling A_plus -> A+ specifically"""
    if epc_score == 'A_plus':
        return 'A+'
    return epc_score.replace('_', '+') if epc_score else 'N/A'

# Load environment variables
load_dotenv()

# FastAPI instance
app = FastAPI(
    title="Azure OpenAI API",
    description="Unified API to interact with Azure OpenAI using FastAPI",
    version="1.0.0"
)

# CORS config
origins = [
    "https://realestate-react-ui-agent.azurewebsites.net",
    "http://localhost:3000"
]

app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["POST", "GET", "OPTIONS"],
    allow_headers=["*"],
)

# === Environment Config ===
AZURE_OPENAI_ENDPOINT = os.getenv("AZURE_OPENAI_ENDPOINT", "").rstrip("/")
AZURE_OPENAI_DEPLOYMENT = os.getenv("AZURE_OPENAI_DEPLOYMENT_NAME")
AZURE_OPENAI_API_VERSION = os.getenv("AZURE_OPENAI_API_VERSION")
AZURE_OPENAI_API_KEY = os.getenv("AZURE_OPENAI_API_KEY")

# Add embedding-specific configuration
AZURE_OPENAI_EMBEDDING_DEPLOYMENT = os.getenv("AZURE_OPENAI_EMBEDDING_DEPLOYMENT", AZURE_OPENAI_DEPLOYMENT)
AZURE_OPENAI_EMBEDDING_MODEL = os.getenv("AZURE_OPENAI_EMBEDDING_MODEL", "text-embedding-ada-002")

# Document processing configuration
UPLOAD_DIR = Path("uploaded_documents")
FAISS_INDEX_DIR = Path("faiss_indexes")
UPLOAD_DIR.mkdir(exist_ok=True)
FAISS_INDEX_DIR.mkdir(exist_ok=True)

# Initialize vector store components with better error handling
def initialize_embeddings():
    """Initialize Azure OpenAI embeddings with proper error handling"""
    try:
        if not AZURE_OPENAI_API_KEY:
            raise ValueError("AZURE_OPENAI_API_KEY not found in environment variables")
        
        if not AZURE_OPENAI_EMBEDDING_DEPLOYMENT:
            raise ValueError("AZURE_OPENAI_EMBEDDING_DEPLOYMENT not found in environment variables")
        
        embeddings = OpenAIEmbeddings(
            deployment=AZURE_OPENAI_EMBEDDING_DEPLOYMENT,
            model=AZURE_OPENAI_EMBEDDING_MODEL,
            openai_api_key=AZURE_OPENAI_API_KEY,
            openai_api_base=AZURE_OPENAI_ENDPOINT,
            openai_api_type="azure",
            openai_api_version=AZURE_OPENAI_API_VERSION,
        )
        
        logger.info(f"✅ Embeddings initialized successfully")
        logger.info(f"   Deployment: {AZURE_OPENAI_EMBEDDING_DEPLOYMENT}")
        logger.info(f"   Model: {AZURE_OPENAI_EMBEDDING_MODEL}")
        logger.info(f"   Endpoint: {AZURE_OPENAI_ENDPOINT}")
        
        return embeddings
        
    except Exception as e:
        logger.error(f"❌ Failed to initialize embeddings: {e}")
        logger.error(f"   Please check your Azure OpenAI configuration in .env file")
        return None

# Initialize embeddings
embeddings = initialize_embeddings()

# Text splitter for document chunking
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=1000,
    chunk_overlap=200,
    length_function=len
)

# In-memory storage for document metadata
document_store: Dict[str, Dict[str, Any]] = {}

# === Schemas ===

class Message(BaseModel):
    role: str
    content: str

class ChatRequest(BaseModel):
     messages: List[Message]

class ChatResponse(BaseModel):
    response: str

class CommentResponse(BaseModel):
    comments: List[str]

class ScoreMeta(BaseModel):
    mae: float
    rmse: float
    r2: float

class FormData(BaseModel):
    region: str
    province: str
    locality: str
    scoreMeta: ScoreMeta

class UserProfile(BaseModel):
    type: str
    objectives: List[str]
    language: str

class CommentRequest(BaseModel):
    formData: FormData
    predictionAll: float
    predictionTop: float
    userProfile: UserProfile

class LLMParamRequest(BaseModel):
    model_name: str

class SuggestionRequest(BaseModel):
    model_name: str
    previous_trials: list

class PropertyFeatures(BaseModel):
    propertyType: str
    subtype: str
    province: str
    locality: str
    postCode: str
    bedroomCount: int
    bathroomCount: int
    toiletCount: int
    roomCount: int
    habitableSurface: int
    facedeCount: int
    buildingConstructionYear: int
    buildingCondition: str
    kitchenType: str
    heatingType: str
    floodZoneType: str
    epcScore: str
    hasLivingRoom: bool
    hasTerrace: bool

class ESGAnalysisRequest(BaseModel):
    propertyFeatures: PropertyFeatures
    estimatedPrice: float
    analysis_depth: str = "detailed"  # "basic", "detailed", "comprehensive"

class ESGAnalysisResponse(BaseModel):
    analysis_points: List[str]
    esg_scores: dict
    recommendations: List[str]
    compliance_status: dict
    financial_impact: dict
    full_report: str

class AgentInsight(BaseModel):
    agent: str
    summary: str

class ESGSummary(BaseModel):
    environment: float
    social: float
    governance: float
    overall: str

class StrategicSummaryRequest(BaseModel):
    price_prediction: float
    esg_summary: ESGSummary
    property_features: PropertyFeatures
    strategic_goals: str  # "invest", "live", "renovate", etc.
    agent_insights: List[AgentInsight]

class StrategicSummaryResponse(BaseModel):
    strategic_positioning: str
    esg_analysis: str
    recommended_actions: str
    clickable_suggestions: List[dict]
    confidence_score: float

# Strategic Analysis Summary - Condensed version for quick overview
class StrategicAnalysisSummaryResponse(BaseModel):
    summary: str
    key_insights: List[str]
    confidence_score: float
    timestamp: str

# RAG Document Upload Models
class DocumentUploadResponse(BaseModel):
    document_id: str
    filename: str
    file_type: str
    size_bytes: int
    chunks_created: int
    status: str
    upload_time: str

class DocumentInfo(BaseModel):
    id: str
    filename: str
    file_type: str
    size_bytes: int
    upload_time: str
    chunks_count: int
    content_preview: str
    tags: List[str]

class DocumentListResponse(BaseModel):
    documents: List[DocumentInfo]
    total_count: int
    total_size_bytes: int

class DocumentQueryRequest(BaseModel):
    query: str
    document_ids: Optional[List[str]] = None
    max_results: int = 5

class DocumentQueryResponse(BaseModel):
    query: str
    results: List[Dict[str, Any]]
    context_used: str
    answer: str
    sources: List[str]

class IndexStatsResponse(BaseModel):
    total_documents: int
    total_chunks: int
    total_size_bytes: int
    index_size_mb: float
    last_updated: str

# === Utility Functions ===

def call_azure_openai_chat(messages: List[dict], temperature: float = 0.7, max_tokens: int = 1500):
    url = f"{AZURE_OPENAI_ENDPOINT}/openai/deployments/{AZURE_OPENAI_DEPLOYMENT}/chat/completions?api-version={AZURE_OPENAI_API_VERSION}"
    
    headers = {
        "Content-Type": "application/json",
        "api-key": AZURE_OPENAI_API_KEY
    }

    payload = {
        "messages": messages,
        "temperature": temperature,
        "max_tokens": max_tokens
    }

    try:
        print("=================================\n")
        print("Payload from UI to to call_azure_openai_cha", payload)  
        response = requests.post(url, headers=headers, json=payload)
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except Exception as e:
        print(f"Azure OpenAI API error: {e}")
        raise HTTPException(status_code=500, detail="Failed to get response from Azure OpenAI.")

# === Document Processing Functions ===

def extract_text_from_pdf(file_path: Path) -> str:
    """Extract text from PDF using PyMuPDF"""
    try:
        doc = fitz.open(file_path)
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path: Path) -> str:
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        logger.error(f"Error extracting text from DOCX {file_path}: {e}")
        return ""

def extract_text_from_txt(file_path: Path) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError:
        # Try with different encoding if UTF-8 fails
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            logger.error(f"Error reading TXT file {file_path}: {e}")
            return ""
    except Exception as e:
        logger.error(f"Error extracting text from TXT {file_path}: {e}")
        return ""

def extract_text_from_file(file_path: Path, file_type: str) -> str:
    """Extract text from various file types"""
    if file_type.lower() == "pdf":
        return extract_text_from_pdf(file_path)
    elif file_type.lower() == "docx":
        return extract_text_from_docx(file_path)
    elif file_type.lower() == "txt":
        return extract_text_from_txt(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")

def generate_document_tags(text: str, filename: str) -> List[str]:
    """Generate tags for a document based on content analysis"""
    tags = []
    
    # File type tag
    if filename.lower().endswith('.pdf'):
        tags.append("PDF")
    elif filename.lower().endswith('.docx'):
        tags.append("Word Document")
    elif filename.lower().endswith('.txt'):
        tags.append("Text File")
    
    # Content-based tags
    text_lower = text.lower()
    
    # ESG and sustainability tags
    if any(keyword in text_lower for keyword in ['esg', 'environment', 'sustainability', 'carbon', 'energy']):
        tags.append("ESG")
    if any(keyword in text_lower for keyword in ['regulation', 'compliance', 'law', 'legal']):
        tags.append("Legal")
    if any(keyword in text_lower for keyword in ['real estate', 'property', 'building', 'construction']):
        tags.append("Real Estate")
    if any(keyword in text_lower for keyword in ['finance', 'investment', 'roi', 'budget']):
        tags.append("Finance")
    if any(keyword in text_lower for keyword in ['report', 'analysis', 'assessment']):
        tags.append("Report")
    
    return tags[:5]  # Limit to 5 tags

def get_faiss_index_path() -> Path:
    """Get the path to the FAISS index"""
    return FAISS_INDEX_DIR / "document_index"

def load_or_create_faiss_index():
    """Load existing FAISS index or create a new one"""
    index_path = get_faiss_index_path()
    
    if index_path.exists():
        try:
            return FAISS.load_local(str(index_path), embeddings)
        except Exception as e:
            logger.warning(f"Could not load existing FAISS index: {e}")
    
    # Create new empty index
    sample_text = ["Initial document for index creation"]
    sample_docs = [Document(page_content=sample_text[0], metadata={"temp": True})]
    vector_store = FAISS.from_documents(sample_docs, embeddings)
    
    # Save the new index
    vector_store.save_local(str(index_path))
    return vector_store

def update_faiss_index(documents: List[Document]):
    """Update FAISS index with new documents"""
    vector_store = load_or_create_faiss_index()
    
    if documents:
        vector_store.add_documents(documents)
        vector_store.save_local(str(get_faiss_index_path()))
    
    return vector_store

# === Routes ===


@app.get("/", tags=["Health"])
def root():
    return {"message": "API LLM V2 is running..."}

@app.get("/health", tags=["Health"])
def health_check():
    """
    Health check endpoint for API connectivity verification.
    Used by frontend to test if the LLM backend service is available.
    """
    return {
        "status": "healthy",
        "service": "LLM Backend API v2",
        "timestamp": datetime.now().isoformat(),
        "port": 8010,
        "endpoints": {
            "chat": "/chat",
            "esg_agent": "/esg_agent", 
            "esg_analysis": "/esg_analysis",
            "documents": "/documents",
            "upload_document": "/upload_document"
        }
    }

@app.post("/chat", response_model=ChatResponse, tags=["Chat"])
def chat(request: ChatRequest):
    if not request.messages:
        raise HTTPException(status_code=400, detail="Missing messages field")

    response_text = call_azure_openai_chat(
        messages=[msg.dict() for msg in request.messages]
    )
    return {"response": response_text}

@app.post("/esg_agent", response_model=ChatResponse, tags=["ESG"])
def esg_agent_chat(request: ChatRequest):
    """
    ESG-focused chat endpoint for the ESG Real Estate Advisor interface.
    Provides specialized responses for ESG compliance, energy performance, and sustainability.
    """
    if not request.messages:
        raise HTTPException(status_code=400, detail="Missing messages field")

    # Add ESG-specific system context to the conversation
    messages = [msg.dict() for msg in request.messages]
    
    # If there's no system message, add one for ESG context
    if not messages or messages[0]["role"] != "system":
        esg_system_message = {
            "role": "system",
            "content": """You are an ESG Real Estate Advisor specializing in sustainable Belgian real estate. 
            You provide expert guidance on:
            - Energy Performance Certificates (EPC) and regulations
            - Belgian grants and subsidies for renovations
            - 2030-2035 compliance planning and execution
            - Sustainable investment strategies
            - Property value enhancement through ESG improvements
            - Regional grant availability in specific Belgian regions
            
            Always provide practical, actionable advice with specific focus on Belgian regulations, 
            EPC improvements, and financial incentives. Use European formatting for amounts (space as thousand separator, € after amount).
            Be concise but comprehensive in your responses."""
        }
        messages.insert(0, esg_system_message)
    
    response_text = call_azure_openai_chat(messages=messages)
    return {"response": response_text}

@app.post("/comment")
def generate_comments(request: CommentRequest):
    form = request.formData
    profile = request.userProfile

    prompt = f"""
    You are a real estate investment advisor for a user profile of type '{profile.type}', 
    with objectives: {', '.join(profile.objectives)}. 
    The property is located in {form.locality}, {form.province}, {form.region}.
    
    Predicted property value: {request.predictionAll:,.0f} € (use European format: space as thousand separator, € after amount)
    Model performance: MAE = {form.scoreMeta.mae:,.0f} €, RMSE = {form.scoreMeta.rmse:,.0f} €, R² = {form.scoreMeta.r2:.2f}

    Focus on:
    1. Regional market trends and investment potential for {form.region}
    2. Property characteristics analysis and investment return strategy
    3. Market insights specific to {form.locality}

    IMPORTANT: When mentioning prices or amounts, always use European format: "417 675 €" (space as thousand separator, € symbol after the amount), NOT "€417,675".

    Generate 2-3 concise, investment-focused comments in English. Each comment should be practical and actionable for real estate investment decisions.
    """

    response_text = call_azure_openai_chat(
        messages=[
            {"role": "system", "content": "You are a helpful real estate AI assistant."},
            {"role": "user", "content": prompt}
        ]
    )

    comments = [c.strip() for c in response_text.strip().split('\n') if c.strip()]
    return {"comments": comments}

from fastapi import APIRouter, HTTPException
import json

router = APIRouter()
app.include_router(router)









@router.post("/suggest-space", tags=["LLM Tuner"])
def suggest_param_space(request: SuggestionRequest):
    model_name = request.model_name
    trials = request.previous_trials

    # Build a summary of previous trials for the prompt
    if not trials:
        trial_summary = "(no prior trials found)"
    else:
        trial_summary = "\n".join([
            f"- Params: {t.get('hyperparameters', t.get('params'))}, "
            f"Score: {t.get('r2_test') or t.get('rmse') or '?'}"
            for t in trials
        ])

    # Construct model-specific prompts
    if model_name.lower() == "catboost":
        prompt = f"""
        You are an expert in Optuna hyperparameter tuning for CatBoost regression models, working on real estate price prediction in production environments.

        Your objective is to generate a high-quality Optuna search space that will help optimize the **Smart Model Ranking Score**, based on the following key criteria:

        - Maximize **R² Test** (ideally ≥ 0.85 for production readiness)
        - Minimize **R² Gap** between Train and Test (ideally ≤ 0.10)
        - Maximize **Generalization Index** (calculated as: 100 - R² Gap × 1000)
        - Minimize **Overfitting Risk** (defined by consistency in train/test metrics)
        - Balance **MAE / RMSE** to ensure stable and reliable error distribution

        ---

        Here are previous tuning trials for '{model_name}':
        {trial_summary}

        Now suggest a full Optuna-compatible JSON search space for **CatBoostRegressor**, ensuring all important parameters are included.

         Prioritize **anti-overfitting and generalization** by:
        - Using lower learning rates
        - Regularization (l2_leaf_reg, random_strength)
        - Limiting depth and overgrown trees
        - Selecting proper grow_policy and bootstrap_type

        ---

        REQUIRED PARAMETERS:
        - learning_rate (suggest_loguniform: 0.01–0.3)
        - depth (suggest_int: 4–10)
        - iterations (suggest_int: 500–2000)
        - l2_leaf_reg (suggest_loguniform: 1.0–10.0)
        - border_count (suggest_int: 32–255)
        - random_strength (suggest_uniform: 0.1–10.0)
        - min_data_in_leaf (suggest_int: 1–20)
        - bootstrap_type (suggest_categorical: ["Bayesian", "Bernoulli", "MVS"])
        - subsample (suggest_uniform: 0.6–1.0) → only if bootstrap_type != "Bayesian"
        - grow_policy (suggest_categorical: ["SymmetricTree", "Depthwise", "Lossguide"])
        - leaf_estimation_method (suggest_categorical: ["Newton", "Gradient"])
        - leaf_estimation_iterations (suggest_int: 1–10)
        - bagging_temperature (suggest_uniform: 0.0–1.0)
        - colsample_bylevel (suggest_uniform: 0.5–1.0)
        - od_type (suggest_categorical: ["IncToDec", "Iter"])
        - od_wait (suggest_int: 10–50)
        - task_type (fixed_value: "CPU")

        ---

        RULES:
        1. Use JSON only (no markdown, no comments)
        2. Use "method": one of suggest_loguniform, suggest_uniform, suggest_int, suggest_categorical, fixed_value
        3. Use "low"/"high" for numeric ranges, "choices" for categorical, and "value" for fixed
        4. Output format must be:

        {{
        "model": "{model_name}",
        "param_space": {{
            "parameter_name": {{
            "method": "suggest_type",
            "low": X,
            "high": Y
            }},
            ...
        }}
        }}

        5. Ensure diversity in trials and robustness across features.
        6. Do NOT propose extreme depths, unbounded iterations, or aggressive learning rates.

        Only output the valid JSON, with no explanations.
        """

    elif model_name.lower() == "xgboost":
        prompt = f"""
        You are an expert in Optuna hyperparameter tuning for XGBoost regression models.

        You are optimizing an **XGBRegressor** for real estate price prediction. Your goal is to maximize a custom Smart Model Ranking Score based on generalization, stability, and production-readiness.

        Key optimization goals:
        - Maximize **R² Test** (≥ 0.85 for production)
        - Minimize **R² Gap** (≤ 0.10 ideal)
        - Maximize **Generalization Index** = 100 - (R² Gap × 1000)
        - Reduce **Overfitting Risk**
        - Ensure consistent and stable **MAE / RMSE** across train/test
        - Avoid overly deep trees or unstable parameter combinations

        Here are previous trials for the model '{model_name}':

        {trial_summary}

        Now suggest a complete Optuna parameter space for XGBRegressor, in valid JSON format, using all relevant parameters.

        REQUIRED PARAMETERS:
        - learning_rate (suggest_float: 0.01–0.3)
        - max_depth (suggest_int: 3–10)
        - min_child_weight (suggest_float: 1.0–10.0)
        - subsample (suggest_float: 0.5–1.0)
        - colsample_bytree (suggest_float: 0.5–1.0)
        - colsample_bylevel (suggest_float: 0.5–1.0)
        - colsample_bynode (suggest_float: 0.5–1.0)
        - gamma (suggest_float: 0.0–5.0)
        - reg_alpha (suggest_float: 0.0–1.0)
        - reg_lambda (suggest_float: 0.0–2.0)
        - n_estimators (suggest_int: 100–2000)
        - max_delta_step (suggest_int: 0–10)
        - grow_policy (suggest_categorical: ["depthwise", "lossguide"])
        - max_leaves (suggest_int: 0–256) → only if grow_policy == "lossguide"
        - tree_method (fixed_value: "auto")

        FORMAT RULES:
        1. Output must be valid JSON only (no markdown, no comments)
        2. Use "method": one of `suggest_float`, `suggest_int`, `suggest_categorical`, `fixed_value`
        3. Use "low"/"high" for ranges, "choices" for categorical values, "value" for fixed
        4. Output must be structured as:

    {{
      "model": "{model_name}",
      "param_space": {{
        "parameter_name": {{
          "method": "suggest_type",
          "low": X,
          "high": Y
        }},
        "categorical_parameter": {{
          "method": "suggest_categorical",
          "choices": [...]
        }},
        "fixed_parameter": {{
          "method": "fixed_value",
          "value": "..."
        }}
      }}
    }}

    Output only valid JSON. No explanation, no extra text.
    """


    else:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported model: {model_name}. Supported models: catboost, xgboost"
        )

    logger.info("=== Prompt received by the API & to be sent to GPT-4.1 ===")
    logger.info(prompt)

    # Call Azure OpenAI chat endpoint
    response_text = call_azure_openai_chat(
        messages=[
            {"role": "system", "content": "You are a helpful assistant specialized in ML hyperparameter tuning."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.2,
        max_tokens=700
    )

    logger.info("=== Response received from GPT-4.1 ===")
    logger.info(response_text)

    cleaned = re.sub(r"^```(?:json)?\s*|\s*```$", "", response_text.strip(), flags=re.IGNORECASE)

    # Attempt to parse JSON response
    try:
        param_space = json.loads(cleaned)
        return {"model": model_name, "param_space": param_space}
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=500,
            detail=f"Invalid JSON returned by LLM:\n{response_text}"
        )



















































# ESG Analysis endpoint
@app.post("/esg_analysis")
async def generate_esg_analysis(request: ESGAnalysisRequest):
    """
    Generate comprehensive ESG analysis for Belgian real estate properties
    using Azure OpenAI with expert knowledge of Belgian regulations and market conditions.
    """
    
    # Extract property features for analysis
    property_data = request.propertyFeatures
    
    # Calculate detailed property characteristics for analysis
    epc_score = property_data.epcScore
    surface = property_data.habitableSurface
    year = property_data.buildingConstructionYear
    locality = property_data.locality
    province = property_data.province
    heating_type = property_data.heatingType
    
    # Calculate energy efficiency metrics
    is_old_building = year < 1980
    is_energy_efficient = epc_score in ['A_plus', 'A', 'B']
    needs_renovation = epc_score in ['E', 'F', 'G']
    
    # Calculate potential savings and renovations
    yearly_energy_cost = surface * (25 if needs_renovation else 15)
    potential_savings = yearly_energy_cost * (0.6 if needs_renovation else 0.3)
    renovation_cost = surface * (250 if needs_renovation else 100)
    
    # Build comprehensive ESG analysis prompt matching the original static analysis quality
    prompt = f"""
    You are a certified ESG and energy performance expert specializing in the Belgian real estate market. Based on the following property characteristics, generate a detailed, structured ESG analysis:

    Property Details:
    - Location: {locality}, {province} ({property_data.postCode})
    - Type: {property_data.propertyType} – {property_data.subtype}
    - Surface: {surface}m²
    - Construction Year: {year}
    - EPC Score: {format_epc_score(epc_score)}
    - Heating System: {heating_type}
    - Condition: {property_data.buildingCondition}
    - Estimated Price: {request.estimatedPrice:,.0f} €

    Your analysis should include **6 well-developed paragraphs**, covering:

    1. **EPC Rating & Energy Performance**  
    Begin with an evaluation of the EPC score and how it compares to national and regional standards. Mention whether the property meets future performance thresholds (e.g. 2030 targets).

    2. **Estimated Energy Consumption & Costs**  
    Estimate primary energy consumption (in kWh/m²/year) based on the EPC score and building size. Provide realistic yearly cost estimates and explain assumptions used.

    3. **Heating System Assessment**  
    Analyze the heating type (“{heating_type}”) in terms of energy efficiency, regulatory compatibility, transition risks, and possible improvements.

    4. **Compliance with Belgian & Regional Regulations**  
    Assess whether the property aligns with current and upcoming rules (e.g. rental bans, renovation mandates in Flanders, Wallonia, Brussels). Focus on relevant provincial context.

    5. **ESG Risks & Opportunities**  
    Identify key ESG risks or weaknesses for this property, and opportunities for improvement (e.g. energy upgrades, accessibility, digital metering, insulation).

    6. **Investment Outlook & Recommendations**  
    Provide actionable suggestions: renovation priorities, estimated investment range (e.g. €{renovation_cost:,.0f}), expected ROI or property value impact (e.g. potential savings of €{potential_savings:,.0f} annually).

    Each paragraph should be 4–6 sentences long, grounded in Belgian legislation and market practices. Use a clear, formal tone suitable for investors or institutional actors. Do not assume any predefined ESG score – derive it from the property characteristics.
"""


    # Call Azure OpenAI for ESG analysis
    response_text = call_azure_openai_chat(
        messages=[
            {"role": "system", "content": "You are a Belgian real estate ESG analysis expert with comprehensive knowledge of environmental, social, and governance factors in Belgian property markets."},
            {"role": "user", "content": prompt}
        ]
    )

    # Parse the response to extract structured ESG data
    # For now, return the full analysis text
    # In the future, could parse specific scores and recommendations
    
    # Parse key analysis points from the response
    analysis_points = []
    recommendations = []
    esg_scores = {"environmental": 7.0, "social": 7.0, "governance": 7.0, "overall": 7.0}
    
    # Extract key points and recommendations from the response
    lines = response_text.split('\n')
    current_section = ""
    
    for line in lines:
        line = line.strip()
        if line.startswith('**') and line.endswith('**'):
            current_section = line.replace('**', '').lower()
        elif line.startswith('- ') and 'recommendation' in current_section:
            recommendations.append(line[2:])
        elif line.startswith('- ') or line.startswith('• '):
            analysis_points.append(line[2:] if line.startswith('- ') else line[2:])
        elif 'score:' in line.lower():
            # Try to extract numerical scores
            try:
                if 'environmental' in line.lower():
                    score = float(line.split(':')[1].split('/')[0].strip())
                    esg_scores["environmental"] = score
                elif 'social' in line.lower():
                    score = float(line.split(':')[1].split('/')[0].strip())
                    esg_scores["social"] = score
                elif 'governance' in line.lower():
                    score = float(line.split(':')[1].split('/')[0].strip())
                    esg_scores["governance"] = score
                elif 'overall' in line.lower():
                    score = float(line.split(':')[1].split('/')[0].strip())
                    esg_scores["overall"] = score
            except (ValueError, IndexError):
                pass  # Keep default scores if parsing fails
    
    # Calculate overall score if not found
    if esg_scores["overall"] == 7.0:
        esg_scores["overall"] = round((esg_scores["environmental"] + esg_scores["social"] + esg_scores["governance"]) / 3, 1)
    
    compliance_status = {
        "energy_compliance": "Compliant" if property_data.epcScore in ["A_plus", "A", "B"] else "Needs Review",
        "building_codes": "Compliant",
        "safety_standards": "Compliant" if property_data.buildingCondition in ["AS_NEW", "GOOD"] else "Needs Assessment"
    }
    
    financial_impact = {
        "energy_cost_annual": f"Estimated {500 + (ord(property_data.epcScore[0]) - ord('A')) * 200} €/year based on EPC {property_data.epcScore}",
        "improvement_cost_estimate": "5,000 - 25,000 € for energy efficiency upgrades",
        "roi_potential": f"ESG improvements could increase property value by {esg_scores['overall'] * 2:.0f}%"
    }
    
    return ESGAnalysisResponse(
        analysis_points=analysis_points[:10] if analysis_points else ["Comprehensive ESG analysis completed", "Property assessment performed"],
        esg_scores=esg_scores,
        recommendations=recommendations[:5] if recommendations else ["Review energy efficiency improvements", "Consider accessibility upgrades"],
        compliance_status=compliance_status,
        financial_impact=financial_impact,
        full_report=response_text
    )


# Quick ESG Analysis endpoint for consistent scoring
@app.post("/esg_quick_analysis")
async def generate_quick_esg_analysis(request: ESGAnalysisRequest):
    """
    Generate quick ESG assessment using Azure OpenAI with consistent scoring
    and brief insights per category.
    """

    # Extract property features
    property_data = request.propertyFeatures
    epc_score = property_data.epcScore
    surface = property_data.habitableSurface
    year = property_data.buildingConstructionYear
    heating_type = property_data.heatingType

    # Prompt for the LLM
    prompt = f"""
    As a Belgian real estate ESG expert, provide a QUICK assessment for this property:

    Property: {property_data.propertyType} in {property_data.locality}, {property_data.province}
    Surface: {surface}m², Year: {year}, EPC: {format_epc_score(epc_score)}, Heating: {heating_type}

    Provide ONLY:
    1. Environmental score (0-10): Based on EPC rating and energy efficiency
    2. Social score (0-10): Based on location and family-friendliness  
    3. Governance score (0-10): Based on building age and condition
    4. Overall ESG score (0-10): Average of the three scores
    5. Three short insights (one per category)

    Format your response EXACTLY like this:
    Environmental Score: X.X/10
    Social Score: X.X/10  
    Governance Score: X.X/10
    Overall ESG Score: X.X/10

    Environmental Insight: [brief insight]
    Social Insight: [brief insight]
    Governance Insight: [brief insight]
    """

    try:
        # Call Azure OpenAI
        response_text = call_azure_openai_chat(
            messages=[
                {"role": "system", "content": "You are a Belgian real estate ESG expert. Provide consistent, professional assessments."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=500
        )

        # Default scores and insights
        esg_scores = {"environmental": 7.0, "social": 7.0, "governance": 7.0, "overall": 7.0}
        insights = {"environment": "", "social": "", "governance": ""}

        # Parse response
        lines = response_text.strip().split('\n')
        for line in lines:
            line = line.strip()
            if "Environmental Score:" in line:
                try:
                    esg_scores["environmental"] = float(line.split(':')[1].split('/')[0].strip())
                except:
                    pass
            elif "Social Score:" in line:
                try:
                    esg_scores["social"] = float(line.split(':')[1].split('/')[0].strip())
                except:
                    pass
            elif "Governance Score:" in line:
                try:
                    esg_scores["governance"] = float(line.split(':')[1].split('/')[0].strip())
                except:
                    pass
            elif "Overall ESG Score:" in line:
                try:
                    esg_scores["overall"] = float(line.split(':')[1].split('/')[0].strip())
                except:
                    pass
            elif "Environmental Insight:" in line:
                insights["environment"] = line.split(":", 1)[1].strip()
            elif "Social Insight:" in line:
                insights["social"] = line.split(":", 1)[1].strip()
            elif "Governance Insight:" in line:
                insights["governance"] = line.split(":", 1)[1].strip()

        # Recalculate overall score if needed
        if esg_scores["overall"] == 7.0:
            esg_scores["overall"] = round(
                (esg_scores["environmental"] + esg_scores["social"] + esg_scores["governance"]) / 3, 1
            )

        # Derive ESG grade
        score = esg_scores["overall"]
        if score >= 8.5:
            grade = "A+"
        elif score >= 7.5:
            grade = "A"
        elif score >= 6.5:
            grade = "B+"
        elif score >= 5.5:
            grade = "B"
        elif score >= 4.5:
            grade = "C"
        else:
            grade = "D"

        return {
            "esg_scores": esg_scores,
            "overall_grade": grade,
            "insights": insights,
            "analysis_summary": response_text,
            "analysis_type": "quick",
            "confidence_level": "high"
        }

    except Exception as e:
        print(f"[✘] Error in quick ESG analysis: {e}")
        raise HTTPException(status_code=500, detail="Failed to generate quick ESG analysis")



# Strategic summary endpoint
@app.post("/strategic_summary", response_model=StrategicSummaryResponse, tags=["ESG"])
async def generate_strategic_summary(request: StrategicSummaryRequest):
    """
    Generate a strategic summary for real estate properties
    focusing on ESG positioning and investment recommendations.
    """
    
    # Extract property features and ESG summary
    property_data = request.property_features
    esg_summary = request.esg_summary
    
    # Build strategic summary prompt
    prompt = f"""
    You are a Belgian real estate expert. Provide a strategic summary for a property based on its features and ESG summary.

    Property Features:
    - Type: {property_data.propertyType} - {property_data.subtype}
    - Location: {property_data.locality}, {property_data.province} ({property_data.postCode})
    - Surface: {property_data.habitableSurface}m²
    - Construction Year: {property_data.buildingConstructionYear}
    - EPC Score: {format_epc_score(property_data.epcScore)}
    - Heating: {property_data.heatingType}
    - Condition: {property_data.buildingCondition}

    ESG Summary:
    - Environmental Score: {esg_summary.environment}
    - Social Score: {esg_summary.social}
    - Governance Score: {esg_summary.governance}
    - Overall ESG Score: {esg_summary.overall}

    Strategic Goals: {request.strategic_goals}

    Provide a strategic positioning statement, ESG analysis, recommended actions, and clickable suggestions in JSON format.

    OUTPUT FORMAT:
    {{
        "strategic_positioning": "string",
        "esg_analysis": "string",
        "recommended_actions": "string",
        "clickable_suggestions": [{{"title": "string", "url": "string"}}],
        "confidence_score": float
    }}

    Only output valid JSON. No markdown, no explanations, no additional text.
    """
    
    logger.info("=== Strategic Summary Prompt ===")
    logger.info(prompt)
    
    # Call Azure OpenAI chat endpoint
    response_text = call_azure_openai_chat(
        messages=[
            {"role": "system", "content": "You are a Belgian real estate expert specializing in ESG analysis and investment strategies."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5,
        max_tokens=800
    )

    logger.info("=== Strategic Summary Response ===")
    logger.info(response_text)
    
    # Attempt to parse JSON response
    try:
        summary_response = json.loads(response_text)
        return summary_response
    except json.JSONDecodeError:
        raise HTTPException(
            status_code=500,
            detail=f"Invalid JSON returned by LLM:\n{response_text}"
        )

# Strategic Summary endpoint for unified ESG analysis
@app.post("/strategic-summary", response_model=StrategicSummaryResponse)
async def create_strategic_summary(request: StrategicSummaryRequest):
    """
    Generate unified strategic summary combining price prediction, ESG analysis, and agent insights
    for consistent user experience across all components.
    """
    
    property_data = request.property_features
    esg = request.esg_summary
    price = request.price_prediction
    
    # Build agent insights summary
    agent_summary = "\n".join([f"- {insight.agent}: {insight.summary}" for insight in request.agent_insights])
    
    # Create comprehensive prompt for strategic analysis
    prompt = f"""
    You are a Belgian real estate strategic advisor. Generate a comprehensive strategic summary that unifies all analysis components.

    PROPERTY OVERVIEW:
    - Location: {property_data.locality}, {property_data.province}
    - Type: {property_data.propertyType} - {property_data.subtype}
    - Predicted Value: €{price:,.0f}
    - EPC Rating: {format_epc_score(property_data.epcScore)}
    - Construction Year: {property_data.buildingConstructionYear}
    - Surface: {property_data.habitableSurface}m²
    - Strategic Goal: {request.strategic_goals}

    ESG PROFILE:
    - Environment Score: {esg.environment}/10
    - Social Score: {esg.social}/10  
    - Governance Score: {esg.governance}/10
    - Overall Rating: {esg.overall}

    AGENT INSIGHTS:
    {agent_summary}

    Generate exactly 4 sections:

    **STRATEGIC POSITIONING:**
    Analyze this property's position in the {property_data.locality} market. Consider the €{price:,.0f} price point, EPC {format_epc_score(property_data.epcScore)} rating, and {property_data.buildingConstructionYear} construction year. How does this compare to local market trends? Include specific market context for {property_data.province}.

    **ESG RISK & OPPORTUNITY ANALYSIS:**
    Based on the {esg.overall} ESG rating (E:{esg.environment}/10, S:{esg.social}/10, G:{esg.governance}/10), identify key risks and opportunities. Focus on Belgian energy regulations, upcoming EPC requirements, and investment implications. Address specific concerns for {format_epc_score(property_data.epcScore)} rated properties.

    **RECOMMENDED ACTIONS:**
    Provide 3-4 specific, actionable recommendations for a "{request.strategic_goals}" strategy. Include timeline, expected costs, and ROI projections. Consider Belgian tax incentives, renovation grants, and regulatory deadlines.

    **NEXT STEPS:**
    Suggest 3 clickable actions with titles and brief explanations that the user can take immediately.

    Keep each section substantial (3-4 sentences) but concise. Use European number formatting (spaces for thousands: €417 675).
    """

    # Call Azure OpenAI for strategic analysis
    response_text = call_azure_openai_chat(
        messages=[
            {"role": "system", "content": "You are a Belgian real estate strategic advisor with expertise in ESG analysis, market positioning, and investment strategies."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,  # Lower temperature for consistent analysis
        max_tokens=1200
    )

    # Parse the response into structured sections
    sections = response_text.split('**')
    strategic_positioning = ""
    esg_analysis = ""
    recommended_actions = ""
    
    for i, section in enumerate(sections):
        section = section.strip()
        if 'STRATEGIC POSITIONING' in section.upper():
            strategic_positioning = sections[i+1].strip() if i+1 < len(sections) else ""
        elif 'ESG RISK' in section.upper():
            esg_analysis = sections[i+1].strip() if i+1 < len(sections) else ""
        elif 'RECOMMENDED ACTIONS' in section.upper():
            recommended_actions = sections[i+1].strip() if i+1 < len(sections) else ""
    
    # Extract clickable suggestions (look for patterns like [Action Title])
    suggestions = []
    suggestion_patterns = re.findall(r'\[([^\]]+)\]', response_text)
    for i, suggestion in enumerate(suggestion_patterns[:3]):
        suggestions.append({
            "id": i + 1,
            "title": suggestion,
            "description": f"Click to explore {suggestion.lower()}",
            "action": suggestion.lower().replace(' ', '_')
        })
    
    # If no suggestions found, provide defaults
    if not suggestions:
        suggestions = [
            {"id": 1, "title": "Compare Market Properties", "description": "Analyze similar properties in the area", "action": "compare_market"},
            {"id": 2, "title": "ESG Improvement Calculator", "description": "Calculate ROI for energy upgrades", "action": "esg_calculator"},
            {"id": 3, "title": "Investment Strategy Guide", "description": "Get personalized investment recommendations", "action": "investment_guide"}
        ]
    
    # Calculate confidence score based on data completeness
    confidence = 0.8  # Base confidence
    if esg.environment > 0 and esg.social > 0 and esg.governance > 0:
        confidence += 0.1
    if len(request.agent_insights) > 2:
        confidence += 0.1
    
    return StrategicSummaryResponse(
        strategic_positioning=strategic_positioning or "Market analysis in progress...",
        esg_analysis=esg_analysis or "ESG assessment completed with current data.",
        recommended_actions=recommended_actions or "Strategic recommendations being prepared...",
        clickable_suggestions=suggestions,
        confidence_score=min(confidence, 1.0)
    )

# Strategic Analysis Summary - Condensed version for quick overview
@app.post("/strategic-analysis-summary", response_model=StrategicAnalysisSummaryResponse)
async def create_strategic_analysis_summary(request: StrategicSummaryRequest):
    """
    Generate a condensed strategic analysis summary combining all data sources
    into a concise overview format for quick decision-making.
    """
    
    property_data = request.property_features
    esg = request.esg_summary
    price = request.price_prediction
    
    # Build agent insights summary
    agent_summary = "\n".join([f"- {insight.agent}: {insight.summary}" for insight in request.agent_insights])
    
    # Create condensed prompt for strategic analysis
    prompt = f"""
    You are a Belgian real estate strategic advisor. Generate a CONDENSED strategic analysis summary for quick decision-making.

    PROPERTY DATA:
    - Location: {property_data.locality}, {property_data.province}
    - Type: {property_data.propertyType} - {property_data.subtype}
    - Predicted Value: €{price:,.0f}
    - EPC Rating: {format_epc_score(property_data.epcScore)}
    - Construction Year: {property_data.buildingConstructionYear}
    - Surface: {property_data.habitableSurface}m²
    - Strategic Goal: {request.strategic_goals}

    ESG PROFILE:
    - Environment Score: {esg.environment}/10
    - Social Score: {esg.social}/10  
    - Governance Score: {esg.governance}/10
    - Overall Rating: {esg.overall}

    AGENT INSIGHTS:
    {agent_summary}

    Generate a CONDENSED analysis in 2-3 sentences that captures the essence of this property's strategic value.
    Focus on: market position, ESG implications, and immediate investment considerations.
    
    Then provide 3-4 key bullet points for quick decision-making.
    
    Format:
    SUMMARY: [2-3 sentences]
    KEY INSIGHTS:
    • [Insight 1]
    • [Insight 2]  
    • [Insight 3]
    • [Insight 4]

    Keep it concise but informative. Use European number formatting (spaces for thousands: €417 675).
    """

    print(">>> ======== DEBUG PROMPT: /strategic-analysis-summary ==============")
    print("/strategic-analysis-summary: ", prompt)

    # Call Azure OpenAI for condensed analysis
    response_text = call_azure_openai_chat(
        messages=[
            {"role": "system", "content": "You are a Belgian real estate strategic advisor specializing in condensed analysis for quick decision-making."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,  # Lower temperature for consistent analysis
        max_tokens=400  # Shorter response for condensed summary
    )

    # Parse the response
    parts = response_text.split('KEY INSIGHTS:')
    summary_text = parts[0].replace('SUMMARY:', '').strip() if len(parts) > 0 else response_text
    
    # Extract key insights
    key_insights = []
    if len(parts) > 1:
        insights_text = parts[1].strip()
        # Split by bullet points and clean up
        insights = insights_text.split('•')
        for insight in insights[1:]:  # Skip first empty element
            cleaned_insight = insight.strip()
            if cleaned_insight:
                key_insights.append(cleaned_insight)
    
    # If no insights found, provide defaults based on data
    if not key_insights:
        key_insights = [
            f"Property valued at €{price:,.0f} in {property_data.locality} market",
            f"ESG rating: {esg.overall} (E:{esg.environment}/10, S:{esg.social}/10, G:{esg.governance}/10)",
            f"EPC {format_epc_score(property_data.epcScore)} rating - consider energy efficiency impact",
            f"Strategic goal: {request.strategic_goals} - aligns with current market conditions"
        ]
    
    # Calculate confidence score based on data completeness
    confidence = 0.8  # Base confidence
    if esg.environment > 0 and esg.social > 0 and esg.governance > 0:
        confidence += 0.1
    if len(request.agent_insights) > 2:
        confidence += 0.1
    
    # Generate timestamp
    from datetime import datetime
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    return StrategicAnalysisSummaryResponse(
        summary=summary_text or "Strategic analysis completed with current market data.",
        key_insights=key_insights,
        confidence_score=min(confidence, 1.0),
        timestamp=timestamp
    )

# === RAG Document Processing Endpoints ===

@app.post("/upload_document", response_model=DocumentUploadResponse, tags=["Documents"])
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form("general")
):
    """
    Upload and process documents for RAG system.
    Supports PDF, DOCX, and TXT files.
    """
    
    # Validate file type
    allowed_types = ["pdf", "docx", "txt"]
    file_extension = file.filename.lower().split('.')[-1] if '.' in file.filename else ""
    
    if file_extension not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type. Allowed types: {', '.join(allowed_types)}"
        )
    
    # Generate unique document ID
    document_id = str(uuid.uuid4())
    timestamp = datetime.now().isoformat()
    
    # Save uploaded file
    file_path = UPLOAD_DIR / f"{document_id}_{file.filename}"
    
    try:
        with open(file_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Extract text from document
        extracted_text = extract_text_from_file(file_path, file_extension)
        
        if not extracted_text.strip():
            raise HTTPException(status_code=400, detail="No text could be extracted from the document")
        
        # Create text chunks
        text_chunks = text_splitter.split_text(extracted_text)
        
        # Create documents for vector store
        documents = []
        for i, chunk in enumerate(text_chunks):
            doc = Document(
                page_content=chunk,
                metadata={
                    "document_id": document_id,
                    "filename": file.filename,
                    "file_type": file_extension,
                    "chunk_index": i,
                    "document_type": document_type,
                    "upload_time": timestamp,
                    "source": f"{file.filename} (chunk {i+1}/{len(text_chunks)})"
                }
            )
            documents.append(doc)
        
        # Update FAISS index
        update_faiss_index(documents)
        
        # Generate tags
        tags = generate_document_tags(extracted_text, file.filename)
        
        # Store document metadata
        document_store[document_id] = {
            "id": document_id,
            "filename": file.filename,
            "file_type": file_extension,
            "size_bytes": file_path.stat().st_size,
            "upload_time": timestamp,
            "chunks_count": len(text_chunks),
            "content_preview": extracted_text[:500] + "..." if len(extracted_text) > 500 else extracted_text,
            "tags": tags,
            "document_type": document_type,
            "file_path": str(file_path)
        }
        
        logger.info(f"Document uploaded successfully: {file.filename} (ID: {document_id})")
        
        return DocumentUploadResponse(
            document_id=document_id,
            filename=file.filename,
            file_type=file_extension,
            size_bytes=file_path.stat().st_size,
            chunks_created=len(text_chunks),
            status="success",
            upload_time=timestamp
        )
        
    except Exception as e:
        # Clean up file if processing failed
        if file_path.exists():
            file_path.unlink()
        
        logger.error(f"Error processing document {file.filename}: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")

@app.get("/documents", response_model=DocumentListResponse, tags=["Documents"])
async def list_documents():
    """Get list of all uploaded documents with metadata"""
    
    documents = []
    total_size = 0
    
    for doc_id, doc_info in document_store.items():
        documents.append(DocumentInfo(**doc_info))
        total_size += doc_info["size_bytes"]
    
    return DocumentListResponse(
        documents=documents,
        total_count=len(documents),
        total_size_bytes=total_size
    )

@app.delete("/documents/{document_id}", tags=["Documents"])
async def delete_document(document_id: str):
    """Delete a document and remove it from the index"""
    
    if document_id not in document_store:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Remove file
        doc_info = document_store[document_id]
        file_path = Path(doc_info["file_path"])
        if file_path.exists():
            file_path.unlink()
        
        # Remove from document store
        del document_store[document_id]
        
        # Note: For production, you would want to rebuild the FAISS index
        # without the deleted document's chunks. For now, we keep the vectors
        # in the index but remove the metadata.
        
        logger.info(f"Document deleted: {document_id}")
        
        return {"status": "success", "message": f"Document {document_id} deleted successfully"}
        
    except Exception as e:
        logger.error(f"Error deleting document {document_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.post("/query_documents", response_model=DocumentQueryResponse, tags=["Documents"])
async def query_documents(request: DocumentQueryRequest):
    """
    Query documents using RAG (Retrieval-Augmented Generation).
    Retrieve relevant document chunks and generate an answer using Azure OpenAI.
    """
    
    try:
        # Load FAISS index
        vector_store = load_or_create_faiss_index()
        
        # Search for relevant documents
        search_results = vector_store.similarity_search_with_score(
            request.query,
            k=request.max_results
        )
        
        # Filter by document IDs if specified
        if request.document_ids:
            search_results = [
                (doc, score) for doc, score in search_results
                if doc.metadata.get("document_id") in request.document_ids
            ]
        
        # Prepare context from retrieved documents
        context_chunks = []
        sources = []
        results = []
        
        for doc, score in search_results:
            context_chunks.append(doc.page_content)
            source = doc.metadata.get("source", "Unknown source")
            if source not in sources:
                sources.append(source)
            
            results.append({
                "content": doc.page_content[:200] + "..." if len(doc.page_content) > 200 else doc.page_content,
                "score": float(score),
                "source": source,
                "metadata": doc.metadata
            })
        
        # Combine context
        context_text = "\n\n".join(context_chunks)
        
        # Generate answer using Azure OpenAI
        if context_chunks:
            prompt = f"""
            Based on the following document excerpts, please answer the user's question comprehensively and accurately.

            Context from documents:
            {context_text}

            User Question: {request.query}

            Please provide a detailed answer based on the information in the documents. If the documents don't contain enough information to fully answer the question, please say so and provide what information is available.

            Answer:
            """
            
            answer = call_azure_openai_chat(
                messages=[
                    {"role": "system", "content": "You are a helpful assistant that answers questions based on provided document context. Always cite the source documents when possible."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,  # Lower temperature for factual responses
                max_tokens=1000
            )
        else:
            answer = "No relevant documents found for your query. Please try a different search term or upload relevant documents."
        
        return DocumentQueryResponse(
            query=request.query,
            results=results,
            context_used=context_text[:1000] + "..." if len(context_text) > 1000 else context_text,
            answer=answer,
            sources=sources
        )
        
    except Exception as e:
        logger.error(f"Error querying documents: {e}")
        raise HTTPException(status_code=500, detail=f"Error querying documents: {str(e)}")

@app.get("/index_stats", response_model=IndexStatsResponse, tags=["Documents"])
async def get_index_stats():
    """Get statistics about the document index"""
    
    total_documents = len(document_store)
    total_chunks = sum(doc["chunks_count"] for doc in document_store.values())
    total_size_bytes = sum(doc["size_bytes"] for doc in document_store.values())
    
    # Calculate index size (approximate)
    index_path = get_faiss_index_path()
    index_size_mb = 0.0
    if index_path.exists():
        try:
            index_files = list(index_path.glob("*"))
            index_size_bytes = sum(f.stat().st_size for f in index_files if f.is_file())
            index_size_mb = index_size_bytes / (1024 * 1024)
        except Exception as e:
            logger.warning(f"Could not calculate index size: {e}")
    
    # Get last update time
    last_updated = "Never"
    if document_store:
        latest_upload = max(doc["upload_time"] for doc in document_store.values())
        last_updated = latest_upload
    
    return IndexStatsResponse(
        total_documents=total_documents,
        total_chunks=total_chunks,
        total_size_bytes=total_size_bytes,
        index_size_mb=round(index_size_mb, 2),
        last_updated=last_updated
    )
app.include_router(router)